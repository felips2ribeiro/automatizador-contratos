import tkinter as tk
from docx import Document
import os


#ADICIONE QUANTAS ENTRADAS QUISER

def preencher_contrato():
    dados_cliente = {
        '{nome}': nome_entry.get(),
        '{cpf}': cpf_entry.get(),
        '{telefone}': telefone_entry.get(),
        '{rua}': rua_entry.get(),
        '{numero}': numero_entry.get(),
        '{bairro}': bairro_entry.get(),
        '{cidade}': cidade_entry.get(),
        '{uf}': uf_entry.get(),
        '{cep}': cep_entry.get(),
        '{email}': email_entry.get(),
        '{modelo}': modelo_entry.get(),
        '{km}': km_entry.get(),
        '{placa}': placa_entry.get(),
        '{cor}': cor_entry.get(),
        '{chassi}': chassi_entry.get(),
        '{renavam}': renavam_entry.get(),
        '{valor}': valor_entry.get(),
        '{forma_de_pagamento}': forma_pagamento_entry.get("1.0", "end-1c"),
	'{data}': data_entry.get()
    }

#COLETANDO OS DADOS DO ARQUIVO DOCX

    doc = Document('contratoeditavel.docx')
    
    for paragrafo in doc.paragraphs:
        for marcador, valor in dados_cliente.items():
            if marcador in paragrafo.text:
                paragrafo.text = paragrafo.text.replace(marcador, valor)
    
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for marcador, valor in dados_cliente.items():
                    if marcador in celula.text:
                        celula.text = celula.text.replace(marcador, valor)

    nome_arquivo = f"CONTRATO VENDA {modelo_entry.get()}{placa_entry.get()}.docx"
    nome_arquivo = nome_arquivo.replace('/', '-')
    pasta_salvar = "coloque o diretorio onde deseja salvar o arquivo"

    caminho_completo = os.path.join(pasta_salvar, nome_arquivo)

    doc.save(caminho_completo)
    status_label.config(text="Contrato preenchido e salvo com sucesso!")
    os.startfile(caminho_completo)

# INICIALIZANDO INTERFACE GRAFICA

root = tk.Tk()
root.title("Preencher Contrato")

# DEFININDO O TAMANHO DA JANELA DA INTERFACE

window_width = 720
window_height = 640
screen_width = root.winfo_screenwidth()
screen_height = root.winfo_screenheight()
x = (screen_width // 2) - (window_width // 2)
y = (screen_height // 2) - (window_height // 2)
root.geometry(f"{window_width}x{window_height}+{x}+{y}")

# CRIANDO AS LABELS, CRIE QUANTAS QUISER DE ACORDO COM AS ENTRADAS

tk.Label(root, text="Nome:").grid(row=0, column=0, sticky='w')
tk.Label(root, text="CPF:").grid(row=1, column=0, sticky='w')
tk.Label(root, text="Telefone:").grid(row=2, column=0, sticky='w')
tk.Label(root, text="Rua:").grid(row=3, column=0, sticky='w')
tk.Label(root, text="Número:").grid(row=4, column=0, sticky='w')
tk.Label(root, text="Bairro:").grid(row=5, column=0, sticky='w')
tk.Label(root, text="Cidade:").grid(row=6, column=0, sticky='w')
tk.Label(root, text="UF:").grid(row=7, column=0, sticky='w')
tk.Label(root, text="CEP:").grid(row=8, column=0, sticky='w')
tk.Label(root, text="E-mail:").grid(row=9, column=0, sticky='w')
tk.Label(root, text="Modelo:").grid(row=10, column=0, sticky='w')
tk.Label(root, text="Quilometragem:").grid(row=11, column=0, sticky='w')
tk.Label(root, text="Placa:").grid(row=12, column=0, sticky='w')
tk.Label(root, text="Cor:").grid(row=13, column=0, sticky='w')
tk.Label(root, text="Chassi:").grid(row=14, column=0, sticky='w')
tk.Label(root, text="Renavam:").grid(row=15, column=0, sticky='w')
tk.Label(root, text="Valor:").grid(row=16, column=0, sticky='w')
tk.Label(root, text="Forma de Pagamento:").grid(row=17, column=0, sticky='w')
tk.Label(root, text="Data:").grid(row=18, column=0, sticky='w')


# CRIANDO AS ENTRADAS NA INTERFACE GRÁFICA

nome_entry = tk.Entry(root, width=40)
nome_entry.grid(row=0, column=1)
cpf_entry = tk.Entry(root, width=40)
cpf_entry.grid(row=1, column=1)
telefone_entry = tk.Entry(root, width=40)
telefone_entry.grid(row=2, column=1)
rua_entry = tk.Entry(root, width=40)
rua_entry.grid(row=3, column=1)
numero_entry = tk.Entry(root, width=40)
numero_entry.grid(row=4, column=1)
bairro_entry = tk.Entry(root, width=40)
bairro_entry.grid(row=5, column=1)
cidade_entry = tk.Entry(root, width=40)
cidade_entry.grid(row=6, column=1)
uf_entry = tk.Entry(root, width=40)
uf_entry.grid(row=7, column=1)
cep_entry = tk.Entry(root, width=40)
cep_entry.grid(row=8, column=1)
email_entry = tk.Entry(root, width=40)
email_entry.grid(row=9, column=1)
modelo_entry = tk.Entry(root, width=40)
modelo_entry.grid(row=10, column=1)
km_entry = tk.Entry(root, width=40)
km_entry.grid(row=11, column=1)
placa_entry = tk.Entry(root, width=40)
placa_entry.grid(row=12, column=1)
cor_entry = tk.Entry(root, width=40)
cor_entry.grid(row=13, column=1)
chassi_entry = tk.Entry(root, width=40)
chassi_entry.grid(row=14, column=1)
renavam_entry = tk.Entry(root, width=40)
renavam_entry.grid(row=15, column=1)
valor_entry = tk.Entry(root, width=40)
valor_entry.grid(row=16, column=1)
forma_pagamento_entry = tk.Text(root, width=40, height=4)
forma_pagamento_entry.grid(row=17, column=1)
data_entry = tk.Entry(root, width=40)
data_entry.grid(row=18, column=1)

# BOTÃO DE GERAR O CONTRATO

preencher_button = tk.Button(root, text="Preencher Contrato", command=preencher_contrato)
preencher_button.grid(row=19, columnspan=2)

# LABEL PARA EXIBIR STATUS

status_label = tk.Label(root, text="")
status_label.grid(row=20, columnspan=2)

root.mainloop()
